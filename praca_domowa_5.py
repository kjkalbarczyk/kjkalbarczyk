#Karolina Kalbarczyk
from docx import Document
document = Document() #tworzenie dokumentu docx
import glob
mylist = [f for f in glob.glob("*.txt")] 
print(mylist)
for f in mylist:
    document.add_heading(f) #dodawanie tytułu
    with open(f) as text: #do odpowiedniego tytulu przypisywanie tekstu z pliku txt
        document.add_paragraph(text)
document.save('moj_word.docx')